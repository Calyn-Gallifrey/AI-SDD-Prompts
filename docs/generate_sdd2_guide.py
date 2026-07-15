#!/usr/bin/env python3
"""Build the canonical UAW-SDD 2.0 operating guide DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
FIGURES = ROOT / "figures"
OUTPUT = ROOT / "UAW-SDD2.0 Skill化方案说明与操作指南.docx"

NAVY = "153B5B"
BLUE = "2F6F9F"
LIGHT_BLUE = "E8F1F8"
GOLD = "B88216"
LIGHT_GOLD = "FFF3D5"
GREEN = "397A55"
LIGHT_GREEN = "E8F4EA"
RED = "A74438"
LIGHT_RED = "FBEAE7"
PURPLE = "6B5795"
LIGHT_PURPLE = "F0ECF8"
INK = "1E2933"
MUTED = "637384"
LINE = "A8BBCB"
PANEL = "F7F9FB"
WHITE = "FFFFFF"
BODY_FONT = "Arial Unicode MS"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LINE, size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:cantSplit")
    tr_pr.append(marker)


def set_repeat_keep(paragraph, *, keep_with_next: bool = False, keep_together: bool = False) -> None:
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.paragraph_format.keep_together = keep_together


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = BODY_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, 8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, value, end])
    tail = paragraph.add_run(" 页")
    set_run_font(tail, 8.5, color=MUTED)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.65)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.75)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    heading_specs = {
        "Title": (26, NAVY, 0, 12),
        "Subtitle": (12, MUTED, 0, 8),
        "Heading 1": (18, NAVY, 16, 8),
        "Heading 2": (13.5, BLUE, 11, 5),
        "Heading 3": (11.5, GREEN, 8, 3),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = style_name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Cm(0.55)
        style.paragraph_format.first_line_indent = Cm(-0.25)
        style.paragraph_format.space_after = Pt(3)

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = BODY_FONT
        code._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        code.font.size = Pt(9)
        code.font.color.rgb = RGBColor.from_string(INK)
        code.paragraph_format.left_indent = Cm(0.35)
        code.paragraph_format.right_indent = Cm(0.35)
        code.paragraph_format.space_before = Pt(4)
        code.paragraph_format.space_after = Pt(6)
        code.paragraph_format.keep_together = True

    header = section.header.paragraphs[0]
    header.text = "UAW-SDD 2.0  |  Skill 化方案说明与操作指南"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        set_run_font(run, 8.5, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    set_repeat_keep(paragraph, keep_with_next=True)


def add_paragraph(doc: Document, text: str = "", *, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    set_repeat_keep(paragraph, keep_together=True)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run_font(run)
        set_repeat_keep(paragraph, keep_together=True)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        run = paragraph.add_run(item)
        set_run_font(run)
        set_repeat_keep(paragraph, keep_together=True)


def add_code(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, PANEL)
    set_cell_border(cell, LINE)
    set_cell_margins(cell, 120, 150, 120, 150)
    paragraph = cell.paragraphs[0]
    paragraph.style = doc.styles["Code Block"]
    for index, line in enumerate(text.splitlines()):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = BODY_FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_callout(doc: Document, title: str, body: str, *, kind: str = "info") -> None:
    palette = {
        "info": (LIGHT_BLUE, BLUE),
        "gate": (LIGHT_GOLD, GOLD),
        "risk": (LIGHT_RED, RED),
        "control": (LIGHT_PURPLE, PURPLE),
        "success": (LIGHT_GREEN, GREEN),
    }
    fill, accent = palette[kind]
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(0.18)
    table.columns[1].width = Cm(16.8)
    left, right = table.rows[0].cells
    set_cell_shading(left, accent)
    set_cell_shading(right, fill)
    for cell in (left, right):
        set_cell_border(cell, fill)
        set_cell_margins(cell, 100, 130, 100, 130)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = right.paragraphs[0]
    first = paragraph.add_run(title + "  ")
    set_run_font(first, 10.5, bold=True, color=accent)
    rest = paragraph.add_run(body)
    set_run_font(rest, 10.5, color=INK)
    set_repeat_keep(paragraph, keep_together=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    widths: list[float] | None = None,
    header_fill: str = NAVY,
    font_size: float = 8.8,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if widths:
        for index, width in enumerate(widths):
            table.columns[index].width = Cm(width)
    header = table.rows[0]
    repeat_table_header(header)
    keep_row_together(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, header_fill)
        set_cell_border(cell, WHITE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(text)
        set_run_font(run, font_size, bold=True, color=WHITE)
    for row_values in rows:
        row = table.add_row()
        keep_row_together(row)
        for index, text in enumerate(row_values):
            cell = row.cells[index]
            set_cell_shading(cell, WHITE if len(table.rows) % 2 else PANEL)
            set_cell_border(cell, LINE)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(text)
            set_run_font(run, font_size, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc: Document, filename: str, caption: str, width: float = 6.75) -> None:
    path = FIGURES / filename
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_after = Pt(8)
    caption_p.paragraph_format.keep_with_next = False
    run = caption_p.add_run(caption)
    set_run_font(run, 8.5, color=MUTED)


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def build_document() -> Document:
    doc = Document()
    configure_document(doc)

    # Cover
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(80)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("UAW-SDD 2.0")
    set_run_font(run, 31, bold=True, color=NAVY)
    title.paragraph_format.space_after = Pt(5)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = subtitle.add_run("Skill 化方案说明与操作指南")
    set_run_font(run, 20, bold=True, color=BLUE)
    subtitle.paragraph_format.space_after = Pt(28)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(16)
    run = rule.add_run("稳定入口  ·  中文主体  ·  确定性状态  ·  可追溯审批  ·  可恢复执行")
    set_run_font(run, 11.5, color=MUTED)
    add_callout(
        doc,
        "核心承诺",
        "开发者仍然只需提交简要提示词并调用 uaw-sdd-ai-coding；状态、审批、Git 范围、失效和恢复全部由 Skill 内部处理，新生成过程资产以简体中文为主体。",
        kind="gate",
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(145)
    metadata = doc.add_table(rows=4, cols=2)
    metadata.autofit = False
    metadata.columns[0].width = Cm(4.2)
    metadata.columns[1].width = Cm(12.5)
    for row, (key, value) in zip(
        metadata.rows,
        [
            ("适用范围", "UAW-SDD 2.0 当前运行基线"),
            ("文档版本", "2.0-hardened / 2026-07-15"),
            ("成熟度目标", "5/5（仓库可控边界）"),
            ("权威规则", "skills/uaw-sdd-ai-coding/references/sdd2-control-contract.md"),
        ],
    ):
        keep_row_together(row)
        for index, text in enumerate((key, value)):
            cell = row.cells[index]
            set_cell_shading(cell, PANEL if index else LIGHT_BLUE)
            set_cell_border(cell, LINE)
            set_cell_margins(cell, 110, 130, 110, 130)
            run = cell.paragraphs[0].add_run(text)
            set_run_font(run, 9.5, bold=index == 0, color=NAVY if index == 0 else INK)

    add_page_break(doc)

    # 1
    add_heading(doc, "1. 使用结论", 1)
    add_callout(
        doc,
        "入口不变",
        "开发者使用方式保持为“简要提示词 + 调用 uaw-sdd-ai-coding”。本次加固没有增加命令、表单、Git 参数或控制文件维护要求。",
        kind="success",
    )
    add_paragraph(
        doc,
        "SDD2.0 由一个总控 Skill、两个下游质量 Skill、九个公开过程资产和一套内部确定性控制记录组成。聊天上下文只承载当前交互，实际恢复点由 Feature 目录中的 .sdd2 控制状态决定。",
    )
    add_bullets(
        doc,
        [
            "总控 uaw-sdd-ai-coding：接收简要需求、编排阶段、生成资产、停在人工 Gate、调度评审和单测。",
            "uaw-code-review：在 SDD 模式下只审查当前冻结范围，输出 Markdown Findings，不修改代码。",
            "uaw-unit-test：先生成或更新真实测试源码；范围重新冻结并复审后，再执行测试并生成 Summary。",
            "控制引擎：绑定 Feature、worktree、分支、Git 基线、审批 revision/hash、范围哈希和归档证据。",
        ],
    )
    add_figure(doc, "sdd2-skill-structure.png", "图 1  SDD2.0 Skill 结构与不变入口")

    add_heading(doc, "2. 开发者入口与最短用法", 1)
    add_paragraph(doc, "开发者只需描述当前需求并点名调用 Skill。以下是建议结构，不是新增强制表单：")
    add_code(
        doc,
        "请使用 uaw-sdd-ai-coding 处理以下需求：\n"
        "功能名称：保单受益人邮箱变更\n"
        "类型：enhancement；模块：policy；Sprint：Sprint8；优先级：P1\n"
        "目标：允许在既有权限和审计规则下变更邮箱\n"
        "变更范围：后端接口、服务、持久化及对应单测\n"
        "禁止变更：数据库表结构、既有接口兼容行为",
    )
    add_paragraph(
        doc,
        "Skill 会从当前提示词和当前代码中安全推导可确认信息，仅针对无法确认且会影响方案的字段提问。模板、历史 Feature 和示例不得补齐业务事实。",
    )
    add_table(
        doc,
        ["字段", "允许值 / 含义", "缺失处理"],
        [
            ["Feature Name", "稳定功能名", "无法确认时询问"],
            ["Feature Type", "query / submit / edit / enhancement / refactor / fix", "从需求语义可安全判断时推导"],
            ["Module / Sprint", "所属模块与迭代", "影响目录定位时必须确认"],
            ["Priority", "P0 / P1 / P2", "无法判断时询问"],
            ["Goal", "一句话行为目标", "必须明确"],
            ["Change Scope", "允许改动的功能与路径边界", "实施前必须明确"],
            ["Forbidden Changes", "禁止触碰的行为或路径", "实施前必须明确"],
        ],
        widths=[3.5, 7.4, 6.0],
    )
    add_callout(
        doc,
        "开发者无需做",
        "无需运行 sdd2_control.py、维护 .sdd2 文件、提供 commit/hash、手工冻结范围或判断恢复点。",
        kind="control",
    )
    add_heading(doc, "2.1 简体中文主体规范", 2)
    add_paragraph(
        doc,
        "三项 Skill 的人类可读运行规则、模板以及新建或修订的九项公开资产，统一以简体中文为主体。唯一语言规则来源是 skills/uaw-sdd-ai-coding/references/language-policy.md。",
    )
    add_bullets(
        doc,
        [
            "标题、说明、事实、判断、风险、验收、评审发现、测试结论和归档总结使用简体中文。",
            "文件名、路径、命令、代码标识符、Schema 键、状态枚举、技术缩写和外部契约原值可保留英文。",
            "用户以英文提交需求时，资产分析和结论仍使用简体中文；需要追溯的原文可按原样引用并标注来源。",
            "控制器在 init 和 record-artifact 时执行语言校验；不符合规范时返回非零并阻塞当前动作。",
            "historical-example 是不可变审计样例，为保持哈希和迁移证据不回写翻译，也不得复制为新资产。",
            "original/ 是保留原文和来源哈希的导入档案，不作为运行规则；面向运行的派生内容必须在 references/ 中以简体中文表达。",
        ],
    )

    # 3-4
    add_heading(doc, "3. 架构与权威层级", 1)
    add_paragraph(doc, "出现冲突时按以下顺序判定；低层规则不得覆盖高层规则：")
    add_numbered(
        doc,
        [
            "SKILL.md：识别用户意图、选择运行模式并维护不变入口。",
            "sdd2-control-contract.md：阶段、状态、审批、失效、范围、恢复和 Archive 资格的唯一事实来源。",
            "sdd2_control.py：确定性执行控制契约；返回非零即硬停止。",
            "references/schemas/：持久化状态、审批和范围的数据结构约束。",
            "language-policy.md：人类可读文件与生成资产的简体中文主体要求。",
            "references/templates/：九个公开资产的人类可读结构。",
            "references/rules/ 与下游 Skill 规则：当前代码实现和测试约束。",
            "examples 与历史 Feature：仅供阅读，不构成需求、审批或运行状态。",
        ],
    )
    add_callout(
        doc,
        "冲突处理",
        "无法按权威层级消解的冲突必须记录为 blocked，并说明恢复条件；禁止模型自行选择看起来合理的解释。",
        kind="risk",
    )

    add_heading(doc, "4. Feature 资产与内部控制", 1)
    add_figure(doc, "sdd2-feature-assets-structure.png", "图 2  九个公开资产、内部控制记录与代码工作树")
    add_table(
        doc,
        ["公开资产", "职责", "上游", "下游 / 完成条件"],
        [
            ["brief-design.md", "固化当前人工简要需求", "当前用户消息", "字段确认后进入 Proposal"],
            ["proposal-input.md", "归一化内部规划输入", "Brief + 定向代码发现", "事实和待确认项明确"],
            ["spec.md", "定义行为与验收", "Proposal + 当前行为", "当前 revision 明确获批"],
            ["design.md", "定义技术增量与约束", "已批 Spec + 当前代码", "当前 revision 明确获批"],
            ["tasks.md", "定义受限 Phase、文件和测试", "已批 Design", "当前 revision 明确获批"],
            ["code-review-findings.md", "记录冻结范围的首轮发现", "冻结 scope + 已批资产", "同一 scope 上记录质量 Gate"],
            ["auto-fix-summary.md", "逐项映射修复或处置", "Findings + 代码变化", "同一复审 scope 上关闭"],
            ["unit-test-summary.md", "记录测试源码与执行证据", "已复审 scope", "测试通过且当前 Summary 获批"],
            ["archive.md", "最终可追溯交付记录", "全部当前资产 + evidence", "预检通过且最终获批"],
        ],
        widths=[4.0, 4.7, 4.3, 4.5],
        font_size=8.2,
    )
    add_paragraph(doc, "内部控制文件位于同一 Feature 的 .sdd2/ 下，由 Skill 自动维护：")
    add_bullets(
        doc,
        [
            "feature-state.json：唯一当前状态和唯一 next_required_action。",
            "gate-approvals.jsonl / events.jsonl：追加式 SHA-256 哈希链。",
            "implementation-scope.json：Git base/head/tree、文件清单和文件哈希。",
            "archive-evidence.json：归档前生成的不可变证据。",
            "revisions/：公开资产的内容寻址 revision 快照。",
        ],
    )

    # 5
    add_heading(doc, "5. 端到端流程与停止点", 1)
    add_figure(doc, "sdd2-end-to-end-flow.png", "图 3  SDD2.0 端到端流程、人工 Gate 与强制回路")
    add_table(
        doc,
        ["阶段", "入口条件", "输出 / 动作", "Gate / 下一阶段条件"],
        [
            ["Brief Design", "当前简要需求", "brief-design.md + 初始化状态", "必填事实确认"],
            ["Proposal", "Brief 已持久化", "proposal-input.md", "无人工 Gate"],
            ["Spec", "Proposal 当前", "spec.md", "停止；新消息批准当前 revision/hash"],
            ["Design", "Spec 当前且已批", "design.md", "停止；新消息批准当前 revision/hash"],
            ["Tasks", "Design 当前且已批", "tasks.md", "停止；新消息批准当前 revision/hash"],
            ["Scope Capture", "Tasks 当前且已批", "干净 Git 基线、路径、Phase、测试路径", "校验通过才实施"],
            ["Implementation", "范围有效", "逐 Phase 生产代码/配置", "每个 Phase 后停止并人工 Review"],
            ["Freeze", "Phase 全部获批", "固定 base/head/tree/files/hash", "范围完整且无漂移"],
            ["Code Review", "当前冻结范围", "Findings + passed/failed/blocked", "必须在同一 scope 通过"],
            ["Auto-fix", "Review 发现或无问题", "Summary + 修复/处置", "代码变化必须 refreeze + full re-review"],
            ["Unit Test Pass 1", "Review/Auto-fix 同 scope 关闭", "生成/更新测试源码", "测试代码变化后重新冻结和完整复审"],
            ["Unit Test Pass 2", "含测试代码的新 scope 已复审", "执行测试 + unit-test-summary.md", "passed 后停止并人工批准 Summary"],
            ["Archive", "全部 Gate 当前有效", "evidence + archive.md", "预检通过；停止并最终人工批准"],
            ["Completed", "Archive 当前 revision 已批", "释放 worktree 锁", "成功终态"],
        ],
        widths=[3.1, 4.5, 5.3, 4.7],
        font_size=7.9,
    )
    add_callout(
        doc,
        "硬停止规则",
        "到达 Spec、Design、Tasks、Unit Test Summary 或 Archive 人工 Gate 后，不得生成下一阶段资产、改代码、调用下游 Skill 或宣称已越过 Gate。",
        kind="gate",
    )

    # 6-7
    add_heading(doc, "6. 人工审批有效性", 1)
    add_paragraph(doc, "有效审批必须同时满足：")
    add_numbered(
        doc,
        [
            "来自 Gate 到达后的新用户消息。",
            "明确指出当前阶段和批准结论。",
            "绑定当前 attempt、artifact revision 和 artifact SHA-256。",
            "以 user-message / human role 记录，并在平台提供时绑定消息 ID。",
            "未被后续内容或范围变化失效。",
        ],
    )
    add_table(
        doc,
        ["可接受示例", "不可接受示例", "原因"],
        [
            ["批准当前 Spec", "继续 / 下一步 / OK", "未明确阶段与批准结论"],
            ["批准当前 Design", "文档中写着 approved", "文件内容不是当前用户审批"],
            ["批准当前 Tasks", "沿用上次批准", "历史消息不绑定当前 revision/hash"],
            ["批准 Phase 2", "模型自审通过", "Phase Review 必须来自人类新消息"],
            ["批准当前 Unit Test Summary", "测试看起来没问题", "未明确批准当前 Summary"],
            ["批准当前 Archive", "示例中的批准文本", "示例不具备审批权"],
        ],
        widths=[5.0, 5.3, 7.0],
        font_size=8.5,
    )
    add_callout(
        doc,
        "模拟模式",
        "真实模式禁止 AI 代替人类审批。演示模式也必须先由当前用户在单独的新消息中明确授权模拟，且授权先被持久化；否则仍然无效。",
        kind="risk",
    )

    add_heading(doc, "7. 状态、终态与恢复", 1)
    add_table(
        doc,
        ["状态", "含义", "允许动作"],
        [
            ["ready", "阶段可开始", "执行唯一 next_required_action"],
            ["executing", "当前动作进行中", "完成并记录当前动作"],
            ["recorded", "资产/结果已记录但尚未到审批态", "执行控制返回的下一动作"],
            ["awaiting-approval", "等待当前普通人工 Gate", "停止，等待明确新消息"],
            ["awaiting-final-approval", "等待 Archive 最终批准", "停止，等待明确新消息"],
            ["blocked", "校验失败或恢复条件未满足", "只修复记录的原因并重试同一动作"],
            ["completed", "成功归档并释放锁", "不可继续修改本 attempt"],
            ["closed-with-risk", "人工接受风险后关闭，非成功交付", "终态；新尝试需明确 restart"],
            ["aborted", "明确终止，非成功交付", "终态；新尝试需明确 restart"],
            ["superseded", "历史示例或被替代尝试", "只读，不可 resume"],
        ],
        widths=[4.1, 6.8, 6.5],
        font_size=8.4,
    )
    add_paragraph(
        doc,
        "跨会话恢复时总控 Skill 先运行内部 resume 校验，并只执行返回的一个 next_required_action。不得根据聊天记忆或 Markdown 状态描述重建进度。跨设备恢复会重新确认仓库身份、分支、锁和文件哈希。",
    )
    add_callout(
        doc,
        "重试边界",
        "blocked 或终态后的新 attempt 必须来自用户新的明确 retry/restart 消息；旧审批、质量结果和 scope 不转移。",
        kind="control",
    )

    # 8-10
    add_heading(doc, "8. 失效矩阵", 1)
    add_paragraph(doc, "发生变化时保留历史，但失效记录不得继续授权流程：")
    add_table(
        doc,
        ["变化", "自动失效范围", "恢复动作"],
        [
            ["Brief / Proposal / Spec", "全部下游审批、Phase Review、scope 与质量 Gate", "从最早受影响阶段重新记录与审批"],
            ["Design", "Design/Tasks 当前审批及全部下游", "重新审批 Design，再生成/审批 Tasks"],
            ["Tasks", "Tasks 审批、Phase Review、scope 与全部质量 Gate", "重新审批 Tasks 并重建实施范围"],
            ["生产 / 测试 / 配置快照", "Code Review、Auto-fix、Unit Test、Summary 审批、Archive evidence", "重新冻结并完整 Review"],
            ["Findings", "Code Review 及后续 Gate", "在当前 scope 重新记录 review 结果"],
            ["Auto-fix Summary", "Auto-fix 及后续 Gate", "重新关闭 Auto-fix"],
            ["Unit Test Summary", "Summary 审批与 Archive", "重新记录测试 Gate 并审批 Summary"],
            ["Archive", "最终 Archive 审批", "重做 Archive check 并重新审批"],
        ],
        widths=[5.0, 7.2, 5.2],
        font_size=8.5,
    )

    add_heading(doc, "9. 实施范围、Git 基线与并行 Feature", 1)
    add_bullets(
        doc,
        [
            "Tasks 获批后，Skill 内部捕获干净基线、当前分支、允许/禁止路径、Phase 和测试路径。",
            "预存未提交改动、detached HEAD、分支漂移、缺失锁、越界路径或过宽通配符都会阻塞实施。",
            "冻结 scope 包含 base/head/tree、精确文件清单、每个文件哈希和整体 snapshot hash。",
            "一个 Git worktree 同时只允许一个 active Feature；并行 Feature 使用独立 worktree。",
            "Feature 路径与仓库身份使用可移植标识，不把本机绝对路径当作跨设备身份。",
        ],
    )
    add_callout(
        doc,
        "范围原则",
        "Code Review 和 Unit Test 只接受总控传入的当前冻结 scope；不得从 git status、上游漂移或 Feature 目录猜测范围。",
        kind="control",
    )

    add_heading(doc, "10. Code Review 与 Auto-fix", 1)
    add_table(
        doc,
        ["模式", "范围", "输出", "合并 Gate 属性"],
        [
            ["SDD_TASK_CODE_REVIEW", "当前 SDD2 冻结 scope", "code-review-findings.md", "正式 SDD 质量 Gate"],
            ["STANDALONE_GIT_RANGE_REVIEW", "固定 base/head + diff hash", "HTML 总结与个人报告", "独立评审"],
            ["STANDALONE_WORKTREE_SNAPSHOT_REVIEW", "固定 HEAD + 路径/文件 hash", "标注 snapshot 的 HTML", "明确非 merge gate"],
        ],
        widths=[5.2, 5.1, 4.5, 3.0],
        font_size=8.3,
    )
    add_paragraph(
        doc,
        "SDD Findings 是当前冻结范围的不可变首轮发现。Auto-fix Summary 必须逐项记录 fixed、accepted、deferred 或 not-applicable 等处置及证据。任何修复造成代码、测试或配置变化后，旧 Review 立即失效，必须重新冻结并执行完整 Code Review；不能只复查已修问题。",
    )

    # 11-12
    add_heading(doc, "11. Unit Test 双阶段流程", 1)
    add_numbered(
        doc,
        [
            "Pass 1：基于已复审 scope、Design 测试策略和目标代码，生成或更新真实测试源码。",
            "测试源码变化后返回总控；重新冻结生产代码 + 测试代码，执行完整 Code Review，并在同一 scope 关闭 Auto-fix。",
            "Pass 2：使用项目真实可执行入口运行窄范围单测，记录命令、环境、退出码、通过/失败/跳过计数、测试文件哈希和 scope hash。",
            "记录 unit-test-summary.md 和 Unit Test Gate；停止等待用户批准当前 Summary。",
        ],
    )
    add_table(
        doc,
        ["结论", "最低证据", "Archive 资格"],
        [
            ["passed", "真实测试入口已执行，退出码与计数证明成功", "可继续，但仍需 Summary 人工批准"],
            ["failed", "命令、失败用例和错误证据", "禁止成功 Archive；修复后重新执行"],
            ["blocked", "缺失依赖/入口/环境及恢复条件", "禁止成功 Archive"],
            ["not-run", "未执行原因和后续命令", "只允许 standalone 记录；SDD 禁止成功 Archive"],
        ],
        widths=[3.2, 9.0, 5.5],
        font_size=8.7,
    )
    add_callout(
        doc,
        "测试真实性",
        "生产代码有可识别目标时，成功 SDD 流程必须至少生成或更新一个匹配已捕获测试路径的测试源码。手工检查可以补充，但不能把 Unit Test Gate 判为 passed。",
        kind="risk",
    )

    add_heading(doc, "12. Archive 与异常关闭", 1)
    add_paragraph(doc, "成功 Archive 必须在同一当前 snapshot 上同时满足：")
    add_bullets(
        doc,
        [
            "Spec、Design、Tasks、Unit Test Summary 当前审批有效；全部 Phase Review 已批准。",
            "Code Review 为 passed；Auto-fix 为 passed 或 not-required；Unit Test 为 passed。",
            "九个公开资产全部存在、已记录且内容哈希未漂移。",
            "Archive evidence 的 Git base/head/tree、scope hash 和文件哈希当前有效。",
            "不存在 blocked、分支漂移、锁不匹配、范围漂移或哈希链损坏。",
        ],
    )
    add_paragraph(
        doc,
        "Archive check 通过后生成并记录 archive.md，然后停止等待最终人工批准。只有当前 Archive revision 获批后，状态才变为 completed 并释放 worktree 锁。",
    )
    add_table(
        doc,
        ["终态", "使用条件", "交付判断"],
        [
            ["completed", "所有当前 Gate 通过且最终 Archive 获批", "成功交付"],
            ["closed-with-risk", "测试失败/阻塞等风险由用户明确接受后关闭", "非成功交付，保留风险"],
            ["aborted", "用户明确终止", "非成功交付"],
            ["superseded", "历史示例或被替代 attempt", "只读，不可恢复为 active"],
        ],
        widths=[4.2, 8.4, 5.0],
        font_size=8.7,
    )

    # 13
    add_heading(doc, "13. 角色与控制时序", 1)
    add_figure(doc, "sdd2-sequence.png", "图 4  开发者、总控、控制引擎、代码仓库与下游 Skill 时序")
    add_paragraph(
        doc,
        "总控 Skill 是唯一对开发者开放的流程入口；控制引擎是机器状态权威；下游 Skill 只处理总控提供的当前冻结范围。角色边界防止 Code Review 自行扩展范围、Unit Test 跳过复审或模型把历史文字当作审批。",
    )

    add_heading(doc, "14. 异常场景与预期控制", 1)
    add_table(
        doc,
        ["场景", "系统行为", "恢复方式"],
        [
            ["跳过 Spec/Design 直接改代码", "控制状态拒绝后续记录并 blocked", "回到最早缺失资产与审批"],
            ["无人工审批继续", "approve 解析拒绝模糊/旧/引用消息", "等待新的阶段明确批准"],
            ["把示例 approved 当审批", "示例为非权威；历史 Feature 为 superseded", "只接受当前用户新消息"],
            ["部分实现后直接 Unit Test", "缺少 Phase Review / Freeze / Review 时阻塞", "按唯一 next action 补齐"],
            ["Auto-fix 后未复审", "scope hash 变化使 Review/Auto-fix 失效", "重新冻结并完整 Review"],
            ["测试源码生成后未复审", "旧 scope 失效，禁止 Pass 2", "冻结含测试源码的新 scope 并完整 Review"],
            ["Unit Test 失败仍 Archive", "Archive check 硬失败", "修复重跑，或显式风险关闭/中止"],
            ["用户中途改需求", "从最早受影响资产自动失效下游", "重新记录、审批和执行受影响阶段"],
            ["会话/设备中断", "resume 校验仓库、锁、分支、哈希并给唯一动作", "只执行返回动作"],
            ["多 Feature 读取错目录", "worktree 锁与 Feature 绑定校验失败", "为并行 Feature 使用独立 worktree"],
            ["重复调用 Skill", "已有状态时 resume；不重复 init", "沿唯一恢复点继续"],
            ["代码变化但文档未同步", "scope/文件哈希与 Gate 不一致，阻止 Archive", "更新受影响资产并重走失效链"],
            ["审批/事件链被改写", "SHA-256 链校验失败，所有写操作硬停止", "人工调查并显式恢复/新 attempt"],
        ],
        widths=[4.7, 7.4, 5.6],
        font_size=7.9,
    )

    # 15-16
    add_heading(doc, "15. 路由、规则与来源边界", 1)
    add_paragraph(
        doc,
        "总控先读取 references/context/routing-index.md，只按当前模块、调用链、模型和测试目标加载最小必要规则。规则选择以当前代码、依赖、附近实现和已批准 Design 为证据，不能仅凭文件名或框架版本猜测。",
    )
    add_table(
        doc,
        ["内容", "权威位置", "使用边界"],
        [
            ["流程控制", "uaw-sdd-ai-coding/references/sdd2-control-contract.md", "唯一流程语义来源"],
            ["资产与交接", "uaw-sdd-ai-coding/references/sdd2-workflow.md", "不重新定义状态/Gate"],
            ["上下文路由", "uaw-sdd-ai-coding/references/context/routing-index.md", "按当前证据选择规则"],
            ["后端/模型规则", "uaw-sdd-ai-coding/references/rules/", "仅在当前代码适用时加载"],
            ["评审模式", "uaw-code-review/references/code-review-rules.md", "三种模式互斥"],
            ["测试路由", "uaw-unit-test/references/testing-profile-routing.md", "依赖和附近测试优先"],
            ["来源映射", "uaw-sdd-ai-coding/references/context/source-provenance.json", "original/ 仅来源追溯，不是运行指令"],
            ["历史 Feature", "sdd2-features/...", "quarantined / superseded，只读样例"],
        ],
        widths=[4.0, 8.6, 5.2],
        font_size=8.2,
    )
    add_callout(
        doc,
        "样例边界",
        "模板和历史 Feature 可帮助理解结构，但不得提供当前业务事实、范围、审批或状态。所有历史 Feature 都应明确标注为 quarantined historical example。",
        kind="risk",
    )

    add_heading(doc, "16. 维护者内部控制", 1)
    add_paragraph(
        doc,
        "以下能力由总控 Skill 内部调用，不构成开发者入口。维护者排障时仍必须以控制契约和命令退出码为准：",
    )
    add_table(
        doc,
        ["内部动作", "目的", "失败规则"],
        [
            ["init / resume", "初始化或验证唯一恢复点", "非零即硬停止"],
            ["record-artifact", "记录 revision/hash 和快照并触发失效", "资产缺失/越序/漂移即停止"],
            ["approve / phase-review", "绑定当前消息、attempt、revision/hash", "模糊/复用/历史消息拒绝"],
            ["capture-scope / freeze-scope", "绑定基线、路径、Git IDs 和文件哈希", "脏基线/越界/分支漂移拒绝"],
            ["quality-gate", "记录 Review、Auto-fix、Unit Test 结果", "必须绑定当前 scope"],
            ["prepare-archive / archive-check", "生成 evidence 并校验成功归档资格", "任一 Gate 失败即拒绝"],
            ["close / restart-attempt", "显式风险关闭、终止或新尝试", "必须有当前用户明确指令"],
        ],
        widths=[5.0, 7.5, 5.3],
        font_size=8.4,
    )

    # 17-18
    add_heading(doc, "17. 验收标准", 1)
    add_table(
        doc,
        ["维度", "5/5 验收条件", "验证方式"],
        [
            ["入口稳定", "开发者仍为简要提示词 + Skill", "静态契约与 Skill 描述检查"],
            ["阶段闭环", "所有阶段有入口、产物、Gate、失败与恢复", "控制状态机单元测试"],
            ["审批可信", "新消息、阶段、attempt、revision/hash 全绑定", "正反例与重放测试"],
            ["范围确定", "干净基线、路径约束、冻结 manifest/hash", "Git 临时仓库场景测试"],
            ["变更失效", "生产/测试/配置变化强制 refreeze + full review", "scope drift / invalidation 测试"],
            ["评审闭环", "Findings、Auto-fix、复审同 scope", "质量 Gate 顺序测试"],
            ["测试真实性", "测试源码 + 可执行证据；失败不可 Archive", "双阶段测试与归档拒绝测试"],
            ["恢复与并行", "唯一 next action；一 worktree 一 active Feature", "resume / lock / restart 测试"],
            ["可追溯归档", "九资产、Git IDs、scope/file hashes 全当前", "archive evidence / check 测试"],
            ["语言一致性", "运行规则和新资产以简体中文为主体，必要英文受控保留", "init / record-artifact 正反例与静态语言校验"],
            ["可维护性", "契约唯一来源、Schema、路由和 provenance 完整", "静态资产验证器"],
        ],
        widths=[3.6, 9.2, 5.1],
        font_size=8.3,
    )
    add_callout(
        doc,
        "平台边界",
        "仓库内可以确定性校验消息记录、哈希链和阶段绑定；平台提供真实 message ID 时同时绑定。平台未提供时使用当前消息文本摘要、时间和哈希链留证，不把无法从仓库证明的 UI 身份事实伪装为已证明。",
        kind="info",
    )

    add_heading(doc, "18. 维护验证命令", 1)
    add_paragraph(doc, "以下命令仅用于 Skill 维护和发布验证，不要求业务开发者执行：")
    add_code(
        doc,
        "python3 -m unittest discover -s skills/uaw-sdd-ai-coding/scripts/tests -v\n"
        "python3 skills/uaw-sdd-ai-coding/scripts/validate_sdd2_assets.py\n"
        "python3 <skill-creator>/scripts/quick_validate.py skills/uaw-sdd-ai-coding\n"
        "python3 <skill-creator>/scripts/quick_validate.py skills/uaw-code-review\n"
        "python3 <skill-creator>/scripts/quick_validate.py skills/uaw-unit-test",
    )
    add_bullets(
        doc,
        [
            "控制单测必须覆盖审批歧义、消息重放、越序资产、脏基线、范围漂移、测试失败、锁、恢复和完整成功路径。",
            "静态验证必须覆盖运行资产、简体中文主体、Schema、引用、来源映射、历史样例隔离和不变入口。",
            "三个 Skill 必须分别通过 Skill 结构验证。",
            "指南和四张图必须重新生成并完成逐页/逐图视觉检查。",
        ],
    )

    add_heading(doc, "19. 发布检查清单", 1)
    add_table(
        doc,
        ["检查项", "通过条件"],
        [
            ["开发者入口", "未引入新命令、参数或手工控制文件"],
            ["控制契约", "与脚本、Schema、Workflow、三 Skill 无冲突"],
            ["九个公开资产", "模板齐全，职责、输入、输出和 Gate 可追溯"],
            ["语言规范", "运行规则及新建或修订资产以简体中文为主体，必要英文不改变机器契约"],
            ["异常路径", "失败/阻塞/重试/风险关闭/中止均有确定行为"],
            ["历史样例", "明确隔离，不含可被误认为当前审批的记录"],
            ["文档与图示", "只描述当前 SDD2.0 运行基线，内容与控制契约一致"],
            ["验证", "单测、静态校验、Skill 校验、文档渲染和 Git diff check 全通过"],
        ],
        widths=[5.0, 12.8],
        font_size=8.8,
    )
    add_callout(
        doc,
        "成熟度结论",
        "当本章全部通过时，SDD2.0 在仓库可控边界达到 5/5：入口稳定、中文资产一致、流程闭环、审批可追溯、范围确定、质量 Gate 可验证、失败可恢复且归档不可伪成功。",
        kind="success",
    )

    return doc


def main() -> None:
    document = build_document()
    temporary = OUTPUT.with_suffix(".tmp.docx")
    document.save(temporary)
    temporary.replace(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
